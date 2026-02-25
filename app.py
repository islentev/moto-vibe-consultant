import streamlit as st
from openai import OpenAI
from duckduckgo_search import DDGS
from docx import Document
from io import BytesIO

# --- ИНИЦИАЛИЗАЦИЯ ИЗ СЕКРЕТОВ ---
# Если файлов secrets нет, используем заглушки для теста
try:
    APP_PASSWORD = st.secrets["APP_PASSWORD"]
    DEEPSEEK_KEY = st.secrets["DEEPSEEK_KEY"]
except:
    APP_PASSWORD = "admin"  # пароль по умолчанию для теста
    DEEPSEEK_KEY = "your_key_here"

client = OpenAI(api_key=DEEPSEEK_KEY, base_url="https://api.deepseek.com")

def get_market_info(query):
    with DDGS() as ddgs:
        results = [r['body'] for r in ddgs.text(f"купить {query} цены рф 2026", max_results=15)]
    return "\n".join(results)

def create_docx(text):
    """Создание Word документа в памяти"""
    doc = Document()
    doc.add_heading('Отчет по подбору мотоцикла', 0)
    doc.add_paragraph(text)
    
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

# --- ИНТЕРФЕЙС ---
st.set_page_config(page_title="MotoVibe Pro", page_icon="🏍️")

with st.sidebar:
    st.title("🔐 Доступ")
    password_input = st.text_input("Введите пароль", type="password")
    
    if password_input != APP_PASSWORD:
        st.error("Доступ ограничен.")
        st.stop()
    
    st.success("Доступ разрешен")
    st.divider()
    
    budget = st.number_input("Бюджет (руб)", min_value=50000, value=600000, step=50000)
    countries = st.multiselect("Страна-производитель", ["Япония", "Китай", "Европа", "США"], default=["Япония", "Китай"])
    selected_countries = ", ".join(countries) if countries else "Любая"
    moto_class = st.selectbox("Класс", ["Naked", "Sport", "Cruiser", "Touring", "Enduro", "Classic"])
    user_height = st.number_input("Ваш рост (см)", 150, 210, 175)
    riding_style = st.radio("Основная локация", ["Только город", "Город + Дача", "Путешествия"])
    city = st.text_input("Город", value="Москва")
    model_count = st.slider("Сколько моделей?", 3, 15, 5)

st.title("🏍️ MotoVibe: Подбор и Анализ Рисков")

# Состояние для хранения ответа ИИ, чтобы кнопка скачивания не исчезала
if 'last_report' not in st.session_state:
    st.session_state.last_report = None

if st.button("Сгенерировать подбор"):
    with st.spinner('Ищем варианты...'):
        search_context = get_market_info(f"{moto_class} за {budget}")
        
        prompt = f"""
        Ты - эксперт-подборщик по мотоциклам. Клиент: рост {user_height}см, бюджет {budget}р, Страны-производители: {selected_countries},
        локация: {riding_style}, город {city}.
        Предложи {model_count} моделей.
        Для каждой: 
        1. Предложи ровно {model_count} моделей, СТРОГО из выбранных стран: {selected_countries}.
        2. Ожидаемый год. 
        3. Риски (Если Япония < 2006 г. или Китай < 2020 г. — детально распиши "возрастные болезни".)
        4. Предупреждение если > 600сс.
        5. Сервис в {city} (Проверь доступность запчастей и профильных сервисов для этих брендов в г. {city} на 2026 год.)
        6. Подходит ли под рост (высота седла vs рост).
        7. Ликвидность: как быстро он его продаст через год.
        8. Технические риски: (твои условия про данные мото).
        9. Юридический совет: на что смотреть в документах этой конкретной модели в 2026 году.
        Вердикт: "Рекомендую к осмотру", "Только для фанатов марки" или "Искать дальше".
        Контекст: {search_context}
        """
        
        try:
            response = client.chat.completions.create(
                model="deepseek-chat",
                messages=[{"role": "user", "content": prompt}]
            )
            st.session_state.last_report = response.choices[0].message.content
        except Exception as e:
            st.error(f"Ошибка API: {e}")

# Если отчет готов, показываем его и даем скачать
if st.session_state.last_report:
    st.markdown("### Результаты анализа")
    st.markdown(st.session_state.last_report)
    
    docx_file = create_docx(st.session_state.last_report)
    
    st.download_button(
        label="📥 Скачать отчет в Word (.docx)",
        data=docx_file,
        file_name=f"moto_selection_{city}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

st.divider()
st.header("💬 Чат с мото-экспертом")

# Инициализируем историю чата, если её еще нет
if "messages" not in st.session_state:
    st.session_state.messages = []

# Отображаем историю сообщений из сессии
for message in st.session_state.messages:
    with st.chat_message(message["role"]):
        st.markdown(message["content"])

# Поле для ввода вопроса пользователем
if prompt_input := st.chat_input("Спроси эксперта (например: 'А почему нет Honda CB400?')"):
    # Добавляем сообщение пользователя в историю
    st.session_state.messages.append({"role": "user", "content": prompt_input})
    with st.chat_message("user"):
        st.markdown(prompt_input)

    # Подготовка контекста для DeepSeek
    # Добавляем основной отчет в начало, если он есть, чтобы ИИ понимал о чем речь
    report_context = f"Контекст предыдущего подбора: {st.session_state.last_report}" if st.session_state.last_report else ""
    
    messages_for_api = [
        {"role": "system", "content": f"Ты — опытный мото-консультант. Отвечай на вопросы пользователя, опираясь на его бюджет и город. {report_context}"}
    ] + st.session_state.messages

    # Запрос к DeepSeek
    with st.chat_message("assistant"):
        try:
            response = client.chat.completions.create(
                model="deepseek-chat",
                messages=messages_for_api,
                stream=False
            )
            full_response = response.choices[0].message.content
            st.markdown(full_response)
            # Сохраняем ответ ИИ в историю
            st.session_state.messages.append({"role": "assistant", "content": full_response})
        except Exception as e:
            st.error(f"Ошибка чата: {e}")
